import os
import tkinter as tk
import tkinter.font as tkFont
import docx
import pdfplumber
from shutil import move
import openpyxl

#-----------------------------------图形界面---------------------------------------------
window = tk.Tk()
window.title('改实验报告')
window.geometry('180x550')

fontStyle = tkFont.Font(family="Lucida Grande", size=20)
name = tk.Label(window,text='aname',fg='red',font=fontStyle)
name.pack()

tk.Label(window,text="封面",bg='lightblue').pack()
coverScore=tk.IntVar()
没对齐=tk.Checkbutton(window,text='没对齐',variable=coverScore, onvalue=1,offvalue=0)
没对齐.pack()
coverScore1=tk.IntVar()
没写日期=tk.Checkbutton(window,text='没写日期',variable=coverScore1, onvalue=1,offvalue=0)
没写日期.pack()

tk.Label(window,text="实验",bg='lightblue').pack()
comment=tk.IntVar()
comment1=tk.IntVar()
comment2=tk.IntVar()
comment3=tk.IntVar()
文字描述简短=tk.Checkbutton(window,text='文字描述简短',variable=comment, onvalue=1,offvalue=0)
文字描述简短.pack()
排版有待提高=tk.Checkbutton(window,text='排版有待提高',variable=comment1, onvalue=1,offvalue=0)
排版有待提高.pack()
缺少截图=tk.Checkbutton(window,text='缺少截图',variable=comment2, onvalue=1,offvalue=0)
缺少截图.pack()
缺少实验步骤=tk.Checkbutton(window,text='缺少实验步骤',variable=comment3, onvalue=1,offvalue=0)
缺少实验步骤.pack()

tk.Label(window,text="总分",bg='lightblue').pack()
expScore=tk.IntVar()
expScore1=tk.IntVar()
expScore2=tk.IntVar()
expScore3=tk.IntVar()
expScore4=tk.IntVar()
expScore5=tk.IntVar()
expScore6=tk.IntVar()
sub=tk.Checkbutton(window,text='-5',variable=expScore, onvalue=1,offvalue=0)
sub.pack()
sub1=tk.Checkbutton(window,text='-5',variable=expScore1, onvalue=1,offvalue=0)
sub1.pack()
sub2=tk.Checkbutton(window,text='-15',variable=expScore2, onvalue=1,offvalue=0)
sub2.pack()
sub3=tk.Checkbutton(window,text='-3',variable=expScore3, onvalue=1,offvalue=0)
sub3.pack()
sub4=tk.Checkbutton(window,text='-2',variable=expScore4, onvalue=1,offvalue=0)
sub4.pack()
sub5=tk.Checkbutton(window,text='-1',variable=expScore5, onvalue=1,offvalue=0)
sub5.pack()
sub6=tk.Checkbutton(window,text='-15',variable=expScore6, onvalue=1,offvalue=0)
sub6.pack()

tk.Label(window,text="完整性",bg='lightblue').pack()
fullScore=tk.IntVar()
没完整=tk.Checkbutton(window,text='没完整',variable=fullScore, onvalue=1,offvalue=0)
没完整.pack()

#-----------------------------------读取文件---------------------------------------------

set = os.listdir(r'C:\Users\86157\Desktop\实验报告')
if 'desktop.ini' in set:
    set.remove('desktop.ini')
for i in set:
    if i.startswith('~$'):
        set.remove(i)
setIter = iter(set)
count = 95 # stop
for i in range(0,count):
    next(setIter)
aname = ''
def open():
    global aname
    nameT=aname
    docStr = next(setIter)
    if docStr.endswith('.docx'):
        file = docx.Document("C:\\Users\\86157\\Desktop\\实验报告\\" + docStr)
        for i in range(20):
            if "报告人" in file.paragraphs[i].text:
                index1 = file.paragraphs[i].text.find('报告人')
                index2 = file.paragraphs[i].text.find('学号')
                aname = file.paragraphs[i].text[index1 + 4:index2].strip()
                name['text'] = aname
                break
    elif docStr.endswith('.pdf'):
        with pdfplumber.open("C:\\Users\\86157\\Desktop\\实验报告\\" + docStr) as pdf:
            paper1 = pdf.pages[0]
            txt = paper1.extract_text()
            a = txt.find('报告人')
            b = txt.find('学号')
            aname = txt[a + 4:b].strip()
            name['text'] = aname
    os.startfile("C:\\Users\\86157\\Desktop\\实验报告\\" + docStr)
    if nameT==aname:
        button['bg']='red'
    global count
    print(count)
    count+=1

workbook = openpyxl.load_workbook(r'C:\Users\86157\Desktop\大学计算机－实验三周四-空白.xlsx')
sheet = workbook.worksheets[0]
def nextone():
    for i in range(2,len(sheet['B'])+1):
        if aname==sheet['B'+str(i)].value:
            sum = 45
            if expScore.get()==1:
                sum-=5
            if expScore1.get()==1:
                sum-=5
            if expScore2.get()==1:
                sum-=15
            if expScore3.get()==1:
                sum-=3
            if expScore4.get()==1:
                sum-=2
            if expScore5.get()==1:
                sum-=1
            if expScore6.get()==1:
                sum-=15
            sheet['C' + str(i)] = sum

            comStr=''
            if comment.get()==1:
                comStr+='文字描述简短，'
            if comment1.get()==1:
                comStr+='排版有待提高，'
            if comment2.get()==1:
                comStr+='缺少截图，'
            if comment3.get()==1:
                comStr+='缺少实验步骤，'
            sheet['D' + str(i)] =comStr

            if coverScore.get()==1 and coverScore1.get()==1:
                sheet['G' + str(i)] =3
                sheet['H'+str(i)] = '没对齐、没写日期'
            elif coverScore.get()==1:
                sheet['G' + str(i)] =4
                sheet['H'+str(i)] = '没对齐'
            elif coverScore1.get()==1 :
                sheet['G' + str(i)] = 4
                sheet['H' + str(i)] = '没写日期'
            else:
                sheet['G' + str(i)] = 5

            if fullScore.get()==1:
                sheet['E' + str(i)] = 4
                sheet['F'+str(i)] = '没写实验结论'
            else:
                sheet['E' + str(i)] = 5

            workbook.save(r'C:\Users\86157\Desktop\大学计算机－实验三周四-空白.xlsx')
            break
    open()
    没对齐.deselect()
    没写日期.deselect()
    文字描述简短.deselect()
    排版有待提高.deselect()
    缺少截图.deselect()
    缺少实验步骤.deselect()
    sub.deselect()
    sub1.deselect()
    sub2.deselect()
    sub3.deselect()
    sub4.deselect()
    sub5.deselect()
    sub6.deselect()
    没完整.deselect()

button = tk.Button(window, text='下一个', width=15, height=2,command=nextone)
button.pack()

open()
window.mainloop()