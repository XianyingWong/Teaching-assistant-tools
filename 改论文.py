import os
import tkinter as tk
import tkinter.font as tkFont
import docx
import pdfplumber
from shutil import move
import openpyxl
import win32clipboard as wc
import win32con

#-----------------------------------图形界面---------------------------------------------
window = tk.Tk()
window.title('改论文')
window.geometry('180x360')

fontStyle = tkFont.Font(family="Lucida Grande", size=20)
name = tk.Label(window,text='aname',fg='red',font=fontStyle)
name.pack()

tk.Label(window,text="论文",bg='lightblue').pack()
comment=tk.IntVar()
comment1=tk.IntVar()
comment2=tk.IntVar()
目录不规范=tk.Checkbutton(window,text='目录不规范',variable=comment, onvalue=1,offvalue=0)
目录不规范.pack()
页眉页脚不规范=tk.Checkbutton(window,text='页眉页脚不规范',variable=comment1, onvalue=1,offvalue=0)
页眉页脚不规范.pack()
标题格式不规范=tk.Checkbutton(window,text='标题格式不规范',variable=comment2, onvalue=1,offvalue=0)
标题格式不规范.pack()

tk.Label(window,text="总分",bg='lightblue').pack()
expScore=tk.IntVar()
expScore1=tk.IntVar()
expScore2=tk.IntVar()
expScore3=tk.IntVar()
expScore4=tk.IntVar()
expScore5=tk.IntVar()
sub=tk.Checkbutton(window,text='-5',variable=expScore, onvalue=1,offvalue=0)
sub.pack()
sub1=tk.Checkbutton(window,text='-5',variable=expScore1, onvalue=1,offvalue=0)
sub1.pack()
sub2=tk.Checkbutton(window,text='-10',variable=expScore2, onvalue=1,offvalue=0)
sub2.pack()
sub3=tk.Checkbutton(window,text='-3',variable=expScore3, onvalue=1,offvalue=0)
sub3.pack()
sub4=tk.Checkbutton(window,text='-2',variable=expScore4, onvalue=1,offvalue=0)
sub4.pack()
sub5=tk.Checkbutton(window,text='-1',variable=expScore5, onvalue=1,offvalue=0)
sub5.pack()

#-----------------------------------读取文件---------------------------------------------

set = os.listdir(r'C:\Users\86157\Desktop\毕业论文')
if 'desktop.ini' in set:
    set.remove('desktop.ini')
for i in set:
    if i.startswith('~$'):
        set.remove(i)
setIter = iter(set)
count = 25 # stop
for i in range(0,count):
    next(setIter)

aname = ''
def open():
    global aname
    nameT=aname
    docStr = next(setIter)
    if docStr.endswith('.docx'):
        file = docx.Document("C:\\Users\\86157\\Desktop\\毕业论文\\" + docStr)
        for i in range(20):
            if "姓名" in file.paragraphs[i].text:
                index = file.paragraphs[i].text.find('姓名')
                aname = file.paragraphs[i].text[index + 3:].strip()
                name['text'] = aname
                break
    elif docStr.endswith('.pdf'):
        with pdfplumber.open("C:\\Users\\86157\\Desktop\\毕业论文\\" + docStr) as pdf:
            paper1 = pdf.pages[0]
            txt = paper1.extract_text()
            a = txt.find('姓名')
            b = txt.find('专业')
            aname = txt[a + 3:b].strip()
            name['text'] = aname
    os.startfile("C:\\Users\\86157\\Desktop\\毕业论文\\" + docStr)
    if nameT==aname:
        button['bg']='red'
    else:
        button['bg'] = 'SystemButtonFace'
    global count
    print(count)
    count+=1

workbook = openpyxl.load_workbook(r'C:\Users\86157\Desktop\大学计算机－实验三周四-空白.xlsx')
sheet = workbook.worksheets[0]
def nextone():
    for i in range(2,len(sheet['B'])+1):
        if aname==sheet['B'+str(i)].value:
            sum = 20
            if expScore.get()==1:
                sum-=5
            if expScore1.get()==1:
                sum-=5
            if expScore2.get()==1:
                sum-=10
            if expScore3.get()==1:
                sum-=3
            if expScore4.get()==1:
                sum-=2
            if expScore5.get()==1:
                sum-=1
            sheet['K' + str(i)] = sum

            comStr=''
            if comment.get()==1:
                comStr+='目录不规范，'
            if comment1.get()==1:
                comStr+='页眉页脚不规范，'
            if comment2.get()==1:
                comStr+='标题格式不规范，'
            sheet['L' + str(i)] =comStr

            workbook.save(r'C:\Users\86157\Desktop\大学计算机－实验三周四-空白.xlsx')
            break
    open()
    目录不规范.deselect()
    页眉页脚不规范.deselect()
    标题格式不规范.deselect()
    sub.deselect()
    sub1.deselect()
    sub2.deselect()
    sub3.deselect()
    sub4.deselect()
    sub5.deselect()

button = tk.Button(window, text='下一个', width=15, height=2,command=nextone)
button.pack()

open()
window.mainloop()